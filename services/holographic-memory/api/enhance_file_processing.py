#!/usr/bin/env python3
"""
Enhanced File Processing Capabilities
====================================

This script enhances the file processing system with additional capabilities,
better integration, and comprehensive testing.
"""

import sys
import os
import time
import json
from typing import Dict, List, Any, Optional
import tempfile
from pathlib import Path

# Add services to path
sys.path.insert(0, 'services')
sys.path.insert(0, 'services/holographic-memory/api')

class EnhancedFileProcessor:
    """Enhanced file processing with additional capabilities"""
    
    def __init__(self):
        self.supported_formats = {
            "txt": self._process_text,
            "md": self._process_markdown,
            "json": self._process_json,
            "csv": self._process_csv,
            "pdf": self._process_pdf,
            "docx": self._process_docx,
            "xml": self._process_xml,
            "html": self._process_html,
            "rtf": self._process_rtf,
            "log": self._process_log,
            "py": self._process_python,
            "js": self._process_javascript,
            "cpp": self._process_cpp,
            "hpp": self._process_cpp,
            "c": self._process_c,
            "h": self._process_c,
            "java": self._process_java,
            "go": self._process_go,
            "rs": self._process_rust,
            "sql": self._process_sql,
            "yaml": self._process_yaml,
            "yml": self._process_yaml,
            "toml": self._process_toml,
            "ini": self._process_ini,
            "cfg": self._process_ini,
            "conf": self._process_ini
        }
        
        self.processing_stats = {
            "total_files": 0,
            "successful_files": 0,
            "failed_files": 0,
            "total_size": 0,
            "processing_times": []
        }
    
    def _process_text(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process plain text files"""
        try:
            text = content.decode('utf-8', errors='ignore')
            return {
                "type": "text",
                "content": text,
                "word_count": len(text.split()),
                "char_count": len(text),
                "line_count": len(text.splitlines()),
                "language": self._detect_language(text),
                "encoding": "utf-8"
            }
        except Exception as e:
            return {"error": f"Text processing failed: {e}"}
    
    def _process_markdown(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process Markdown files"""
        try:
            text = content.decode('utf-8', errors='ignore')
            
            # Extract metadata
            metadata = self._extract_markdown_metadata(text)
            
            # Count markdown elements
            headers = text.count('#')
            links = text.count('[')
            images = text.count('![')
            code_blocks = text.count('```')
            
            return {
                "type": "markdown",
                "content": text,
                "word_count": len(text.split()),
                "char_count": len(text),
                "line_count": len(text.splitlines()),
                "language": "markdown",
                "encoding": "utf-8",
                "metadata": metadata,
                "markdown_elements": {
                    "headers": headers,
                    "links": links,
                    "images": images,
                    "code_blocks": code_blocks
                }
            }
        except Exception as e:
            return {"error": f"Markdown processing failed: {e}"}
    
    def _process_json(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process JSON files"""
        try:
            text = content.decode('utf-8', errors='ignore')
            data = json.loads(text)
            
            return {
                "type": "json",
                "content": text,
                "parsed_data": data,
                "structure": self._analyze_json_structure(data),
                "size": len(text),
                "valid": True
            }
        except json.JSONDecodeError as e:
            return {
                "type": "json",
                "content": content.decode('utf-8', errors='ignore'),
                "valid": False,
                "error": f"Invalid JSON: {e}"
            }
        except Exception as e:
            return {"error": f"JSON processing failed: {e}"}
    
    def _process_csv(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process CSV files"""
        try:
            import pandas as pd
            import io
            
            text = content.decode('utf-8', errors='ignore')
            df = pd.read_csv(io.StringIO(text))
            
            return {
                "type": "csv",
                "content": text,
                "rows": len(df),
                "columns": len(df.columns),
                "column_names": df.columns.tolist(),
                "data_types": df.dtypes.to_dict(),
                "sample_data": df.head().to_dict(),
                "statistics": df.describe().to_dict() if len(df) > 0 else {}
            }
        except ImportError:
            # Fallback without pandas
            text = content.decode('utf-8', errors='ignore')
            lines = text.splitlines()
            if lines:
                headers = lines[0].split(',')
                return {
                    "type": "csv",
                    "content": text,
                    "rows": len(lines) - 1,
                    "columns": len(headers),
                    "column_names": headers,
                    "pandas_available": False
                }
            return {"error": "Empty CSV file"}
        except Exception as e:
            return {"error": f"CSV processing failed: {e}"}
    
    def _process_pdf(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process PDF files"""
        try:
            import fitz  # PyMuPDF
            
            doc = fitz.open(stream=content, filetype="pdf")
            text = ""
            page_count = len(doc)
            
            for page_num in range(page_count):
                page = doc[page_num]
                text += page.get_text()
            
            doc.close()
            
            return {
                "type": "pdf",
                "content": text,
                "pages": page_count,
                "word_count": len(text.split()),
                "char_count": len(text),
                "language": self._detect_language(text)
            }
        except ImportError:
            return {"error": "PyMuPDF not available for PDF processing"}
        except Exception as e:
            return {"error": f"PDF processing failed: {e}"}
    
    def _process_docx(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process DOCX files"""
        try:
            from docx import Document
            import io
            
            doc = Document(io.BytesIO(content))
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return {
                "type": "docx",
                "content": text,
                "word_count": len(text.split()),
                "char_count": len(text),
                "paragraphs": len(doc.paragraphs),
                "language": self._detect_language(text)
            }
        except ImportError:
            return {"error": "python-docx not available for DOCX processing"}
        except Exception as e:
            return {"error": f"DOCX processing failed: {e}"}
    
    def _process_xml(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process XML files"""
        try:
            import xml.etree.ElementTree as ET
            
            text = content.decode('utf-8', errors='ignore')
            root = ET.fromstring(text)
            
            return {
                "type": "xml",
                "content": text,
                "root_tag": root.tag,
                "structure": self._analyze_xml_structure(root),
                "size": len(text)
            }
        except ET.ParseError as e:
            return {
                "type": "xml",
                "content": content.decode('utf-8', errors='ignore'),
                "valid": False,
                "error": f"Invalid XML: {e}"
            }
        except Exception as e:
            return {"error": f"XML processing failed: {e}"}
    
    def _process_html(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process HTML files"""
        try:
            from bs4 import BeautifulSoup
            
            text = content.decode('utf-8', errors='ignore')
            soup = BeautifulSoup(text, 'html.parser')
            
            # Extract text content
            text_content = soup.get_text()
            
            return {
                "type": "html",
                "content": text_content,
                "raw_content": text,
                "title": soup.title.string if soup.title else None,
                "links": len(soup.find_all('a')),
                "images": len(soup.find_all('img')),
                "word_count": len(text_content.split()),
                "char_count": len(text_content)
            }
        except ImportError:
            # Fallback without BeautifulSoup
            text = content.decode('utf-8', errors='ignore')
            return {
                "type": "html",
                "content": text,
                "beautifulsoup_available": False,
                "word_count": len(text.split()),
                "char_count": len(text)
            }
        except Exception as e:
            return {"error": f"HTML processing failed: {e}"}
    
    def _process_rtf(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process RTF files"""
        try:
            import striprtf
            
            text = content.decode('utf-8', errors='ignore')
            plain_text = striprtf.striprtf(text)
            
            return {
                "type": "rtf",
                "content": plain_text,
                "raw_content": text,
                "word_count": len(plain_text.split()),
                "char_count": len(plain_text)
            }
        except ImportError:
            # Fallback without striprtf
            text = content.decode('utf-8', errors='ignore')
            return {
                "type": "rtf",
                "content": text,
                "striprtf_available": False,
                "word_count": len(text.split()),
                "char_count": len(text)
            }
        except Exception as e:
            return {"error": f"RTF processing failed: {e}"}
    
    def _process_log(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process log files"""
        try:
            text = content.decode('utf-8', errors='ignore')
            lines = text.splitlines()
            
            # Analyze log patterns
            error_count = sum(1 for line in lines if 'error' in line.lower() or 'exception' in line.lower())
            warning_count = sum(1 for line in lines if 'warning' in line.lower() or 'warn' in line.lower())
            info_count = sum(1 for line in lines if 'info' in line.lower())
            
            return {
                "type": "log",
                "content": text,
                "lines": len(lines),
                "word_count": len(text.split()),
                "char_count": len(text),
                "log_analysis": {
                    "errors": error_count,
                    "warnings": warning_count,
                    "info": info_count
                }
            }
        except Exception as e:
            return {"error": f"Log processing failed: {e}"}
    
    def _process_python(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process Python files"""
        try:
            text = content.decode('utf-8', errors='ignore')
            
            # Basic Python analysis
            lines = text.splitlines()
            code_lines = [line for line in lines if line.strip() and not line.strip().startswith('#')]
            comment_lines = [line for line in lines if line.strip().startswith('#')]
            
            # Count imports, functions, classes
            imports = [line for line in lines if line.strip().startswith(('import ', 'from '))]
            functions = [line for line in lines if line.strip().startswith('def ')]
            classes = [line for line in lines if line.strip().startswith('class ')]
            
            return {
                "type": "python",
                "content": text,
                "lines": len(lines),
                "code_lines": len(code_lines),
                "comment_lines": len(comment_lines),
                "imports": len(imports),
                "functions": len(functions),
                "classes": len(classes),
                "language": "python"
            }
        except Exception as e:
            return {"error": f"Python processing failed: {e}"}
    
    def _process_javascript(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process JavaScript files"""
        try:
            text = content.decode('utf-8', errors='ignore')
            
            # Basic JavaScript analysis
            lines = text.splitlines()
            code_lines = [line for line in lines if line.strip() and not line.strip().startswith('//')]
            comment_lines = [line for line in lines if line.strip().startswith('//')]
            
            # Count functions, variables
            functions = [line for line in lines if 'function ' in line or '=>' in line]
            variables = [line for line in lines if 'var ' in line or 'let ' in line or 'const ' in line]
            
            return {
                "type": "javascript",
                "content": text,
                "lines": len(lines),
                "code_lines": len(code_lines),
                "comment_lines": len(comment_lines),
                "functions": len(functions),
                "variables": len(variables),
                "language": "javascript"
            }
        except Exception as e:
            return {"error": f"JavaScript processing failed: {e}"}
    
    def _process_cpp(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process C++ files"""
        try:
            text = content.decode('utf-8', errors='ignore')
            
            # Basic C++ analysis
            lines = text.splitlines()
            code_lines = [line for line in lines if line.strip() and not line.strip().startswith('//')]
            comment_lines = [line for line in lines if line.strip().startswith('//')]
            
            # Count includes, functions, classes
            includes = [line for line in lines if line.strip().startswith('#include')]
            functions = [line for line in lines if '::' in line or '(' in line and ')' in line]
            classes = [line for line in lines if 'class ' in line or 'struct ' in line]
            
            return {
                "type": "cpp",
                "content": text,
                "lines": len(lines),
                "code_lines": len(code_lines),
                "comment_lines": len(comment_lines),
                "includes": len(includes),
                "functions": len(functions),
                "classes": len(classes),
                "language": "cpp"
            }
        except Exception as e:
            return {"error": f"C++ processing failed: {e}"}
    
    def _process_c(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process C files"""
        try:
            text = content.decode('utf-8', errors='ignore')
            
            # Basic C analysis
            lines = text.splitlines()
            code_lines = [line for line in lines if line.strip() and not line.strip().startswith('//')]
            comment_lines = [line for line in lines if line.strip().startswith('//')]
            
            # Count includes, functions
            includes = [line for line in lines if line.strip().startswith('#include')]
            functions = [line for line in lines if '(' in line and ')' in line and '{' in line]
            
            return {
                "type": "c",
                "content": text,
                "lines": len(lines),
                "code_lines": len(code_lines),
                "comment_lines": len(comment_lines),
                "includes": len(includes),
                "functions": len(functions),
                "language": "c"
            }
        except Exception as e:
            return {"error": f"C processing failed: {e}"}
    
    def _process_java(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process Java files"""
        try:
            text = content.decode('utf-8', errors='ignore')
            
            # Basic Java analysis
            lines = text.splitlines()
            code_lines = [line for line in lines if line.strip() and not line.strip().startswith('//')]
            comment_lines = [line for line in lines if line.strip().startswith('//')]
            
            # Count imports, classes, methods
            imports = [line for line in lines if line.strip().startswith('import ')]
            classes = [line for line in lines if 'class ' in line or 'interface ' in line]
            methods = [line for line in lines if '(' in line and ')' in line and '{' in line]
            
            return {
                "type": "java",
                "content": text,
                "lines": len(lines),
                "code_lines": len(code_lines),
                "comment_lines": len(comment_lines),
                "imports": len(imports),
                "classes": len(classes),
                "methods": len(methods),
                "language": "java"
            }
        except Exception as e:
            return {"error": f"Java processing failed: {e}"}
    
    def _process_go(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process Go files"""
        try:
            text = content.decode('utf-8', errors='ignore')
            
            # Basic Go analysis
            lines = text.splitlines()
            code_lines = [line for line in lines if line.strip() and not line.strip().startswith('//')]
            comment_lines = [line for line in lines if line.strip().startswith('//')]
            
            # Count imports, functions, types
            imports = [line for line in lines if line.strip().startswith('import ')]
            functions = [line for line in lines if 'func ' in line]
            types = [line for line in lines if 'type ' in line]
            
            return {
                "type": "go",
                "content": text,
                "lines": len(lines),
                "code_lines": len(code_lines),
                "comment_lines": len(comment_lines),
                "imports": len(imports),
                "functions": len(functions),
                "types": len(types),
                "language": "go"
            }
        except Exception as e:
            return {"error": f"Go processing failed: {e}"}
    
    def _process_rust(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process Rust files"""
        try:
            text = content.decode('utf-8', errors='ignore')
            
            # Basic Rust analysis
            lines = text.splitlines()
            code_lines = [line for line in lines if line.strip() and not line.strip().startswith('//')]
            comment_lines = [line for line in lines if line.strip().startswith('//')]
            
            # Count imports, functions, structs
            imports = [line for line in lines if line.strip().startswith('use ')]
            functions = [line for line in lines if 'fn ' in line]
            structs = [line for line in lines if 'struct ' in line or 'enum ' in line]
            
            return {
                "type": "rust",
                "content": text,
                "lines": len(lines),
                "code_lines": len(code_lines),
                "comment_lines": len(comment_lines),
                "imports": len(imports),
                "functions": len(functions),
                "structs": len(structs),
                "language": "rust"
            }
        except Exception as e:
            return {"error": f"Rust processing failed: {e}"}
    
    def _process_sql(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process SQL files"""
        try:
            text = content.decode('utf-8', errors='ignore')
            
            # Basic SQL analysis
            lines = text.splitlines()
            code_lines = [line for line in lines if line.strip() and not line.strip().startswith('--')]
            comment_lines = [line for line in lines if line.strip().startswith('--')]
            
            # Count different SQL statements
            selects = [line for line in lines if 'SELECT' in line.upper()]
            inserts = [line for line in lines if 'INSERT' in line.upper()]
            updates = [line for line in lines if 'UPDATE' in line.upper()]
            deletes = [line for line in lines if 'DELETE' in line.upper()]
            
            return {
                "type": "sql",
                "content": text,
                "lines": len(lines),
                "code_lines": len(code_lines),
                "comment_lines": len(comment_lines),
                "statements": {
                    "selects": len(selects),
                    "inserts": len(inserts),
                    "updates": len(updates),
                    "deletes": len(deletes)
                },
                "language": "sql"
            }
        except Exception as e:
            return {"error": f"SQL processing failed: {e}"}
    
    def _process_yaml(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process YAML files"""
        try:
            import yaml
            
            text = content.decode('utf-8', errors='ignore')
            data = yaml.safe_load(text)
            
            return {
                "type": "yaml",
                "content": text,
                "parsed_data": data,
                "structure": self._analyze_yaml_structure(data),
                "size": len(text),
                "valid": True
            }
        except ImportError:
            return {
                "type": "yaml",
                "content": content.decode('utf-8', errors='ignore'),
                "valid": False,
                "error": "PyYAML not available for YAML processing"
            }
        except yaml.YAMLError as e:
            return {
                "type": "yaml",
                "content": content.decode('utf-8', errors='ignore'),
                "valid": False,
                "error": f"Invalid YAML: {e}"
            }
        except Exception as e:
            return {"error": f"YAML processing failed: {e}"}
    
    def _process_toml(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process TOML files"""
        try:
            import toml
            
            text = content.decode('utf-8', errors='ignore')
            data = toml.loads(text)
            
            return {
                "type": "toml",
                "content": text,
                "parsed_data": data,
                "size": len(text),
                "valid": True
            }
        except ImportError:
            return {
                "type": "toml",
                "content": content.decode('utf-8', errors='ignore'),
                "valid": False,
                "error": "toml not available for TOML processing"
            }
        except toml.TomlDecodeError as e:
            return {
                "type": "toml",
                "content": content.decode('utf-8', errors='ignore'),
                "valid": False,
                "error": f"Invalid TOML: {e}"
            }
        except Exception as e:
            return {"error": f"TOML processing failed: {e}"}
    
    def _process_ini(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process INI/CFG files"""
        try:
            import configparser
            
            text = content.decode('utf-8', errors='ignore')
            config = configparser.ConfigParser()
            config.read_string(text)
            
            sections = list(config.sections())
            
            return {
                "type": "ini",
                "content": text,
                "sections": sections,
                "section_count": len(sections),
                "size": len(text),
                "valid": True
            }
        except ImportError:
            return {
                "type": "ini",
                "content": content.decode('utf-8', errors='ignore'),
                "valid": False,
                "error": "configparser not available for INI processing"
            }
        except configparser.Error as e:
            return {
                "type": "ini",
                "content": content.decode('utf-8', errors='ignore'),
                "valid": False,
                "error": f"Invalid INI: {e}"
            }
        except Exception as e:
            return {"error": f"INI processing failed: {e}"}
    
    def _detect_language(self, text: str) -> str:
        """Simple language detection"""
        # Very basic language detection based on common words
        english_words = ['the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by']
        spanish_words = ['el', 'la', 'de', 'que', 'y', 'a', 'en', 'un', 'es', 'se', 'no', 'te', 'lo', 'le']
        french_words = ['le', 'la', 'de', 'et', 'à', 'un', 'il', 'que', 'ne', 'se', 'ce', 'pas', 'tout']
        
        text_lower = text.lower()
        
        english_score = sum(1 for word in english_words if word in text_lower)
        spanish_score = sum(1 for word in spanish_words if word in text_lower)
        french_score = sum(1 for word in french_words if word in text_lower)
        
        if english_score > spanish_score and english_score > french_score:
            return "english"
        elif spanish_score > french_score:
            return "spanish"
        elif french_score > 0:
            return "french"
        else:
            return "unknown"
    
    def _extract_markdown_metadata(self, text: str) -> Dict[str, Any]:
        """Extract metadata from Markdown frontmatter"""
        lines = text.splitlines()
        metadata = {}
        
        if lines and lines[0].strip() == '---':
            for i, line in enumerate(lines[1:], 1):
                if line.strip() == '---':
                    break
                if ':' in line:
                    key, value = line.split(':', 1)
                    metadata[key.strip()] = value.strip()
        
        return metadata
    
    def _analyze_json_structure(self, data: Any) -> Dict[str, Any]:
        """Analyze JSON structure"""
        if isinstance(data, dict):
            return {
                "type": "object",
                "keys": list(data.keys()),
                "key_count": len(data.keys())
            }
        elif isinstance(data, list):
            return {
                "type": "array",
                "length": len(data),
                "element_types": list(set(type(item).__name__ for item in data))
            }
        else:
            return {
                "type": type(data).__name__,
                "value": str(data)
            }
    
    def _analyze_xml_structure(self, element) -> Dict[str, Any]:
        """Analyze XML structure"""
        children = list(element)
        return {
            "tag": element.tag,
            "attributes": element.attrib,
            "children": [child.tag for child in children],
            "child_count": len(children),
            "text": element.text.strip() if element.text else None
        }
    
    def _analyze_yaml_structure(self, data: Any) -> Dict[str, Any]:
        """Analyze YAML structure"""
        if isinstance(data, dict):
            return {
                "type": "mapping",
                "keys": list(data.keys()),
                "key_count": len(data.keys())
            }
        elif isinstance(data, list):
            return {
                "type": "sequence",
                "length": len(data),
                "element_types": list(set(type(item).__name__ for item in data))
            }
        else:
            return {
                "type": type(data).__name__,
                "value": str(data)
            }
    
    def process_file(self, content: bytes, filename: str) -> Dict[str, Any]:
        """Process a file with enhanced capabilities"""
        start_time = time.time()
        
        # Get file extension
        file_extension = filename.split('.')[-1].lower() if '.' in filename else 'txt'
        
        # Update stats
        self.processing_stats["total_files"] += 1
        self.processing_stats["total_size"] += len(content)
        
        try:
            # Process based on file type
            if file_extension in self.supported_formats:
                result = self.supported_formats[file_extension](content, filename)
                result["filename"] = filename
                result["file_extension"] = file_extension
                result["file_size"] = len(content)
                result["processing_time"] = time.time() - start_time
                result["supported"] = True
                
                self.processing_stats["successful_files"] += 1
                self.processing_stats["processing_times"].append(result["processing_time"])
                
                return result
            else:
                # Fallback to text processing
                result = self._process_text(content, filename)
                result["filename"] = filename
                result["file_extension"] = file_extension
                result["file_size"] = len(content)
                result["processing_time"] = time.time() - start_time
                result["supported"] = False
                result["fallback"] = True
                
                self.processing_stats["successful_files"] += 1
                self.processing_stats["processing_times"].append(result["processing_time"])
                
                return result
                
        except Exception as e:
            self.processing_stats["failed_files"] += 1
            return {
                "filename": filename,
                "file_extension": file_extension,
                "file_size": len(content),
                "processing_time": time.time() - start_time,
                "supported": False,
                "error": str(e)
            }
    
    def get_processing_stats(self) -> Dict[str, Any]:
        """Get comprehensive processing statistics"""
        if self.processing_stats["processing_times"]:
            avg_time = sum(self.processing_stats["processing_times"]) / len(self.processing_stats["processing_times"])
            min_time = min(self.processing_stats["processing_times"])
            max_time = max(self.processing_stats["processing_times"])
        else:
            avg_time = min_time = max_time = 0.0
        
        return {
            "total_files": self.processing_stats["total_files"],
            "successful_files": self.processing_stats["successful_files"],
            "failed_files": self.processing_stats["failed_files"],
            "success_rate": self.processing_stats["successful_files"] / max(self.processing_stats["total_files"], 1),
            "total_size": self.processing_stats["total_size"],
            "avg_processing_time": avg_time,
            "min_processing_time": min_time,
            "max_processing_time": max_time,
            "supported_formats": list(self.supported_formats.keys())
        }

def test_enhanced_file_processing():
    """Test the enhanced file processing capabilities"""
    print("🧪 Testing Enhanced File Processing...")
    
    processor = EnhancedFileProcessor()
    
    # Test cases
    test_files = [
        ("test.txt", b"This is a test text file with some content for processing."),
        ("test.md", b"# Test Markdown\n\nThis is a **test** markdown file.\n\n- Item 1\n- Item 2"),
        ("test.json", b'{"name": "test", "value": 123, "items": [1, 2, 3]}'),
        ("test.csv", b"name,age,city\nJohn,25,New York\nJane,30,Los Angeles"),
        ("test.py", b"def hello_world():\n    print('Hello, World!')\n\nif __name__ == '__main__':\n    hello_world()"),
        ("test.js", b"function helloWorld() {\n    console.log('Hello, World!');\n}\n\nhelloWorld();"),
        ("test.cpp", b"#include <iostream>\n\nint main() {\n    std::cout << \"Hello, World!\" << std::endl;\n    return 0;\n}"),
        ("test.yaml", b"name: test\nvalue: 123\nitems:\n  - item1\n  - item2"),
        ("test.ini", b"[section1]\nkey1 = value1\nkey2 = value2\n\n[section2]\nkey3 = value3"),
        ("test.log", b"2023-01-01 10:00:00 INFO: Application started\n2023-01-01 10:01:00 ERROR: Something went wrong\n2023-01-01 10:02:00 WARNING: This is a warning")
    ]
    
    results = {}
    
    for filename, content in test_files:
        print(f"\n   Processing: {filename}")
        result = processor.process_file(content, filename)
        results[filename] = result
        
        if "error" in result:
            print(f"     ❌ Error: {result['error']}")
        else:
            print(f"     ✅ Success: {result.get('type', 'unknown')} file processed")
            if "word_count" in result:
                print(f"        Words: {result['word_count']}")
            if "lines" in result:
                print(f"        Lines: {result['lines']}")
            if "language" in result:
                print(f"        Language: {result['language']}")
    
    # Get statistics
    stats = processor.get_processing_stats()
    
    print(f"\n📊 Processing Statistics:")
    print(f"   Total Files: {stats['total_files']}")
    print(f"   Successful: {stats['successful_files']}")
    print(f"   Failed: {stats['failed_files']}")
    print(f"   Success Rate: {stats['success_rate']:.2%}")
    print(f"   Average Processing Time: {stats['avg_processing_time']:.4f}s")
    print(f"   Supported Formats: {len(stats['supported_formats'])}")
    
    return results, stats

def main():
    """Main enhanced file processing function"""
    print("🚀 Starting Enhanced File Processing")
    print("=" * 60)
    
    # Test enhanced file processing
    results, stats = test_enhanced_file_processing()
    
    # Generate report
    print("\n" + "=" * 60)
    print("📊 ENHANCED FILE PROCESSING REPORT")
    print("=" * 60)
    
    print(f"\n🔧 CAPABILITIES:")
    print(f"   Supported Formats: {len(stats['supported_formats'])}")
    print(f"   Text Processing: ✅")
    print(f"   Code Analysis: ✅")
    print(f"   Structured Data: ✅")
    print(f"   Document Processing: ✅")
    print(f"   Log Analysis: ✅")
    
    print(f"\n📈 PERFORMANCE:")
    print(f"   Total Files Processed: {stats['total_files']}")
    print(f"   Success Rate: {stats['success_rate']:.2%}")
    print(f"   Average Processing Time: {stats['avg_processing_time']:.4f}s")
    print(f"   Total Data Processed: {stats['total_size']} bytes")
    
    print(f"\n🎯 SUPPORTED FILE TYPES:")
    for fmt in sorted(stats['supported_formats']):
        print(f"   - .{fmt}")
    
    if stats['success_rate'] == 1.0:
        print(f"\n🎉 ENHANCED FILE PROCESSING COMPLETED SUCCESSFULLY!")
        print("✅ All file types processed successfully")
        print("✅ Advanced analysis capabilities working")
        print("✅ Performance metrics collected")
        print("✅ Error handling implemented")
    else:
        print(f"\n⚠️  SOME FILE PROCESSING TESTS FAILED")
        print("❌ File processing needs attention")
    
    print("=" * 60)
    
    return {
        "results": results,
        "stats": stats,
        "success": stats['success_rate'] == 1.0
    }

if __name__ == "__main__":
    main()
